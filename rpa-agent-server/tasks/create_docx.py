"""
create_docx.py — 使用 python-docx 生成标准 .docx 文件

输入 taskParams:
  deliverableId : str        — 对应 LTC-OS deliverables 表的 ID
  title         : str        — 文档标题（投标文件/合同意见/交付包等）
  documentType  : str (可选) — "tender" | "contract" | "handover" | "proposal"（影响页眉/封面）
  companyName   : str (可选) — 公司名称（出现在封面和页眉）
  sections      : list[dict] — 章节内容
    每个 section:
      heading : str   — 章节标题（一级标题）
      body    : str   — 正文（Markdown 子集：**粗体**, - bullet, ## 二级标题）
      level   : int (可选) — 标题层级 1-3，默认 1

输出:
  { fileUrl: "/files/xxx.docx" }
"""

import os
import re
import uuid
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


STORAGE_DIR = os.path.join(os.path.dirname(__file__), "..", "storage")

COLOR_PRIMARY = RGBColor(0x1A, 0x3A, 0x6B)
COLOR_ACCENT  = RGBColor(0x00, 0x8B, 0xD8)


def _set_heading_style(paragraph, level: int = 1):
    """设置标题段落颜色"""
    for run in paragraph.runs:
        run.font.color.rgb = COLOR_PRIMARY if level == 1 else COLOR_ACCENT


def _add_cover_page(doc: Document, title: str, document_type: str, company_name: str):
    """添加封面"""
    # 封面标题
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    # 文档类型标签
    type_labels = {
        "tender": "投标文件",
        "contract": "合同审查意见",
        "handover": "项目交付包",
        "proposal": "方案建议书",
    }
    type_label = type_labels.get(document_type, "技术文档")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(type_label)
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = COLOR_ACCENT

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 公司名
    if company_name:
        comp = doc.add_paragraph()
        comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        comp_run = comp.add_run(company_name)
        comp_run.font.size = Pt(14)
        comp_run.font.bold = True

    # 日期
    from datetime import date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(date.today().strftime("%Y 年 %m 月 %d 日")).font.size = Pt(12)

    # 分页
    doc.add_page_break()


def _parse_body_text(doc: Document, body: str):
    """
    解析简化 Markdown 并添加段落
    支持：
      ## 二级标题
      - / * bullet 项
      **粗体** 内联
      普通文本段落
    """
    lines = body.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()  # 空行保留间距
            continue

        # 二级标题
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            _set_heading_style(h, level=2)
            continue

        # 三级标题
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            _set_heading_style(h, level=3)
            continue

        # Bullet
        if stripped.startswith(("- ", "* ", "• ")):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(para, stripped[2:])
            continue

        # 普通段落（支持内联粗体）
        para = doc.add_paragraph()
        _add_inline_formatting(para, stripped)


def _add_inline_formatting(para, text: str):
    """处理 **粗体** 内联标记"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


async def create_docx_task(task_params: dict[str, Any]) -> dict[str, Any]:
    os.makedirs(STORAGE_DIR, exist_ok=True)

    deliverable_id = task_params.get("deliverableId", str(uuid.uuid4()))
    title = task_params.get("title", "技术文档")
    document_type = task_params.get("documentType", "proposal")
    company_name = task_params.get("companyName", "")
    sections: list[dict] = task_params.get("sections", [])

    doc = Document()

    # --- 全局样式 ---
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    # 页边距
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # 封面
    _add_cover_page(doc, title, document_type, company_name)

    # 目录占位（标记）
    toc_placeholder = doc.add_paragraph("[目录]")
    toc_placeholder.runs[0].italic = True
    toc_placeholder.runs[0].font.color.rgb = COLOR_ACCENT
    doc.add_page_break()

    # 正文章节
    for section_data in sections:
        heading_text = section_data.get("heading", "")
        body_text    = section_data.get("body", "")
        level        = section_data.get("level", 1)

        if heading_text:
            h = doc.add_heading(heading_text, level=min(level, 3))
            _set_heading_style(h, level=level)

        if body_text:
            _parse_body_text(doc, body_text)

        doc.add_paragraph()  # 章节间空行

    # 保存
    filename = f"{deliverable_id}.docx"
    filepath = os.path.join(STORAGE_DIR, filename)
    doc.save(filepath)

    return {"fileUrl": f"/files/{filename}"}
